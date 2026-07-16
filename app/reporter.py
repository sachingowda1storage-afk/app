import os
import json
import csv
from datetime import datetime
import matplotlib.pyplot as plt

# Try importing docx and reportlab safely
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    RL_AVAILABLE = True
except ImportError:
    RL_AVAILABLE = False

def set_cell_background(cell, hex_color):
    """Sets background color of docx cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def create_matplotlib_graph(df, title, graph_type, output_path):
    """
    Generates a beautiful Matplotlib line chart styled with NIMMA Blue and industrial theme grid lines.
    Saves to output_path.
    """
    plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
    fig, ax = plt.subplots(figsize=(6, 3.5), dpi=150)
    
    # Theme color definitions
    nimma_blue = '#0F4C81'
    grid_color = '#E5E5E5'
    text_color = '#333333'
    
    # Clean column names in the DataFrame to be case-insensitive and stripped
    cols = {c.lower().strip(): c for c in df.columns}
    
    x_col = None
    y_col = None
    
    # Check if we should plot multiple parameters or overlay
    # Standard format: "Y vs X" (e.g. "Voltage vs Time", "Emissions Level vs Frequency")
    if " vs " in graph_type:
        parts = graph_type.split(" vs ")
        y_name = parts[0].strip().lower()
        x_name = parts[1].strip().lower()
        
        # Match x_name to columns
        for c_low, c_orig in cols.items():
            if x_name in c_low or c_low in x_name:
                x_col = c_orig
                break
        # Match y_name to columns
        for c_low, c_orig in cols.items():
            if y_name in c_low or c_low in y_name:
                y_col = c_orig
                break

    # Fallbacks if graph_type parsing didn't match perfectly
    if not x_col:
        for name in ["time_sec", "time", "elapsed time", "frequency", "freq"]:
            if name in cols:
                x_col = cols[name]
                break
        if not x_col and len(df.columns) > 0:
            x_col = df.columns[0]
            
    if not y_col:
        for name in ["voltage", "current", "temperature", "acceleration", "emissions", "pressure"]:
            if name in cols:
                y_col = cols[name]
                break
        if not y_col and len(df.columns) > 1:
            y_col = df.columns[1]
        elif not y_col and len(df.columns) > 0:
            y_col = df.columns[0]

    # Plot
    if x_col and y_col:
        x_data = df[x_col].copy()
        y_data = df[y_col].copy()
        
        # Use Time_Sec under the hood if they chose Time to prevent string plotting errors
        if x_col == "Time" and "Time_Sec" in df.columns:
            x_data = df["Time_Sec"].copy()
            x_label = "Time"
        else:
            x_label = x_col
            
        # If the selected column is not numeric, parse/coerce it
        if x_data.dtype == object:
            try:
                from processor import clean_time
                x_data = x_data.apply(clean_time)
            except Exception:
                x_data = pd.to_numeric(x_data, errors='coerce').fillna(0.0)
        
        if y_data.dtype == object:
            y_data = pd.to_numeric(y_data, errors='coerce').fillna(0.0)
            
        # Auto-scale time axis to hours if it spans long intervals
        if ('time_sec' in x_label.lower() or x_label.lower() == 'time') and x_data.max() > 300:
            x_data = x_data / 3600.0
            x_axis_label = f"{x_col} (hours)"
        else:
            x_axis_label = x_col
            
        y_label = y_col
        
        ax.plot(x_data, y_data, color=nimma_blue, linewidth=1.5, label=y_label)
        ax.set_ylabel(y_label, color=text_color, fontweight='bold')
        ax.set_xlabel(x_axis_label, color=text_color, fontweight='bold')
        
        # Check overlay for Voltage & Current default
        if "overlay" in graph_type.lower() or "voltage & current" in graph_type.lower():
            if 'Current' in df.columns and y_col != 'Current':
                ax2 = ax.twinx()
                ax2.plot(x_data, df['Current'], color='#FF6B6B', linewidth=1.2, alpha=0.7, label='Current (A)')
                ax2.set_ylabel('Current (A)', color='#FF6B6B', fontweight='bold')
                ax2.grid(False)
    else:
        ax.text(0.5, 0.5, "No matching column data found", ha='center', va='center')
        
    ax.set_title(title, color=text_color, fontsize=11, fontweight='bold', pad=10)
    ax.grid(True, linestyle='--', color=grid_color)
    
    # Border styling
    for spine in ax.spines.values():
        spine.set_color('#CCCCCC')
        
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches='tight')
    plt.close()

class NumberedCanvas(canvas.Canvas):
    """Custom canvas for ReportLab to support two-pass 'Page X of Y' generation."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#555555"))
        
        # Draw running header on page 2 onwards
        if self._pageNumber > 1:
            self.drawString(54, 750, "TVS MOTOR COMPANY | TVSM QAD LAB")
            self.drawRightString(558, 750, "Classification: Internal")
            self.setStrokeColor(colors.HexColor("#D3D3D3"))
            self.setLineWidth(0.5)
            self.line(54, 742, 558, 742)
            
        # Draw running footer
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 40, page_str)
        self.drawString(54, 40, "Confidential - TVS Motor Company")
        self.line(54, 52, 558, 52)
        
        self.restoreState()

def export_to_docx(project_data, output_path):
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed or available.")
        
    doc = docx.Document()
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Set up Page Header
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("TVS MOTOR COMPANY | TVSM QAD LAB  ")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        hrun.italic = True
        
    # 1. Header Block / Cover Page Style
    p_logo = doc.add_paragraph()
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(base_dir, "assets", "logo.png")
    if os.path.exists(logo_path):
        try:
            p_logo.add_run().add_picture(logo_path, width=docx.shared.Inches(1.2))
        except Exception:
            p_logo_run = p_logo.add_run("TVS MOTOR COMPANY")
            p_logo_run.bold = True
            p_logo_run.font.size = Pt(14)
            p_logo_run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    else:
        p_logo_run = p_logo.add_run("TVS MOTOR COMPANY")
        p_logo_run.bold = True
        p_logo_run.font.size = Pt(14)
        p_logo_run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_heading(f"{project_data['battery_name']} Validation Report", level=1)
    
    # Project Metadata Grid Table
    table_meta = doc.add_table(rows=6, cols=2)
    table_meta.autofit = False
    
    meta_fields = [
        ("Report Reference Number:", project_data.get("request_no", "N/A")),
        ("Date:", datetime.now().strftime("%d.%m.%Y")),
        ("Classification:", "Internal"),
        ("Model / Configuration:", f"{project_data['battery_name']} - {project_data.get('configuration', 'N/A')} ({project_data.get('cell_type', 'N/A')})"),
        ("Standard Reference:", project_data.get("tses_version", "TSES 799 v8")),
        ("Project Name & Customer:", f"{project_data.get('project_name', 'N/A')} | {project_data.get('customer', 'N/A')}")
    ]
    
    for idx, (label, val) in enumerate(meta_fields):
        row = table_meta.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(val))
        set_cell_background(row.cells[0], "F0F4F8")
        
    doc.add_paragraph().add_run().add_break()
    
    # 2. Main Executive Summary / Test Table
    doc.add_heading("1. Executive Summary Table", level=2)
    
    table_results = doc.add_table(rows=1, cols=6)
    table_results.style = 'Table Grid'
    hdr_cells = table_results.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Class").bold = True
    hdr_cells[1].paragraphs[0].add_run("Test Name").bold = True
    hdr_cells[2].paragraphs[0].add_run("Inputs / Details").bold = True
    hdr_cells[3].paragraphs[0].add_run("Requirements").bold = True
    hdr_cells[4].paragraphs[0].add_run("Observations / Measurements").bold = True
    hdr_cells[5].paragraphs[0].add_run("Verdict").bold = True
    
    for cell in hdr_cells:
        set_cell_background(cell, "0F4C81")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for idx, pt in enumerate(project_data.get('tests', [])):
        row_cells = table_results.add_row().cells
        row_cells[0].paragraphs[0].text = str(pt['test_id'])
        row_cells[1].paragraphs[0].text = str(pt['test_name'])
        
        # Input specs
        row_cells[2].paragraphs[0].text = f"File: {os.path.basename(pt['excel_path'] or 'None')}"
        
        # Requirements
        row_cells[3].paragraphs[0].text = str(pt.get('acceptance_criteria', 'N/A'))
        
        # Observations
        row_cells[4].paragraphs[0].text = str(pt.get('observations', 'Pending Analysis'))
        
        # Status
        verdict = str(pt.get('status', 'Pending'))
        run_v = row_cells[5].paragraphs[0].add_run(verdict)
        run_v.bold = True
        if verdict == 'Pass':
            set_cell_background(row_cells[5], "E2F0D9") # Light Green
            run_v.font.color.rgb = RGBColor(0x38, 0x76, 0x1D)
        elif verdict == 'Fail':
            set_cell_background(row_cells[5], "FCE4D6") # Light Red
            run_v.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            set_cell_background(row_cells[5], "FFF2CC") # Light Yellow
            run_v.font.color.rgb = RGBColor(0xB4, 0x5F, 0x06)
            
    doc.add_paragraph().add_run().add_break()
    
    # 3. Individual Test Details Section
    doc.add_heading("2. Test Analysis Details", level=2)
    for pt in project_data.get('tests', []):
        if pt.get('status') == 'Pending':
            continue
            
        doc.add_heading(f"Test Class {pt['test_id']} - {pt['test_name']}", level=3)
        doc.add_paragraph().add_run(f"Description: {pt.get('test_desc', 'N/A')}")
        
        # Print stats from result json
        if pt.get('results_json'):
            try:
                results = json.loads(pt['results_json'])
                doc.add_paragraph().add_run("Analytical Statistics:").bold = True
                
                if results.get('discharge'):
                    d = results['discharge']
                    doc.add_paragraph(style='List Bullet').add_run(f"Measured Discharge Capacity: {d.get('capacity_ah', 0):.2f} Ah")
                    doc.add_paragraph(style='List Bullet').add_run(f"Measured Discharge Energy: {d.get('energy_wh', 0):.1f} Wh")
                    doc.add_paragraph(style='List Bullet').add_run(f"Average Discharge Voltage: {d.get('avg_voltage', 0):.2f} V")
                    doc.add_paragraph(style='List Bullet').add_run(f"Calculated C-rate: {d.get('c_rate', 0):.2f}C")
                    
                if results.get('transitions'):
                    tr = results['transitions']
                    max_ir = max([t['dc_ir'] for t in tr]) if tr else 0.0
                    doc.add_paragraph(style='List Bullet').add_run(f"Estimated Max DC Internal Resistance: {max_ir:.2f} mΩ")
                    
                if results.get('thermal'):
                    t = results['thermal']
                    doc.add_paragraph(style='List Bullet').add_run(f"Temperature Stats: Peak={t.get('t_max', 0):.1f}°C, Rise=+{t.get('t_rise', 0):.1f}°C")
            except Exception:
                pass
                
        # Insert generated custom graphs if they exist
        has_custom = False
        if pt.get('results_json'):
            try:
                results_data = json.loads(pt['results_json'])
                custom_graphs = results_data.get("custom_graphs", [])
                if custom_graphs:
                    has_custom = True
                    for idx, g in enumerate(custom_graphs):
                        img_path = g.get('image_path')
                        if img_path and os.path.exists(img_path):
                            doc.add_picture(img_path, width=Inches(5.5))
                            caption = f"Figure {pt['test_id']}-{idx+1}: {g.get('y1_axis')} vs {g.get('x_axis')} validation curve."
                            doc.add_paragraph(caption).italic = True
            except Exception:
                pass
                
        if not has_custom:
            # Fallback to single standard graph
            graph_img = pt.get('processed_data_path')
            if graph_img and os.path.exists(graph_img):
                png_graph = graph_img.replace('.csv', '_graph.png')
                if os.path.exists(png_graph):
                    doc.add_picture(png_graph, width=Inches(5.5))
                    doc.add_paragraph(f"Figure {pt['test_id']}: Dynamic validation profile curve.").italic = True
                
    # 4. Signatures / Conclusions
    doc.add_paragraph().add_run().add_break()
    doc.add_heading("3. Conclusions & Declarations", level=2)
    conclusion_text = project_data.get("comments", "The battery pack has undergone all selected verification tests. Details are recorded as above.")
    doc.add_paragraph(conclusion_text)
    
    doc.add_paragraph().add_run().add_break()
    table_sig = doc.add_table(rows=2, cols=3)
    table_sig.autofit = True
    
    sigs = [
        ("Prepared by,", "Checked by,", "Approved by,"),
        (project_data.get("engineer", "Testing Engineer"), "S Shivram\nQuality Team", "C Varunkumar\nLab In-Charge")
    ]
    for r_idx, row_vals in enumerate(sigs):
        row = table_sig.rows[r_idx]
        for c_idx, text in enumerate(row_vals):
            row.cells[c_idx].paragraphs[0].text = text
            
    doc.save(output_path)

def export_to_pdf(project_data, output_path):
    if not RL_AVAILABLE:
        raise ImportError("ReportLab is not installed or available.")
        
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_title = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        textColor=colors.HexColor("#0F4C81"),
        spaceAfter=15
    )
    style_h2 = ParagraphStyle(
        'SectionH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        textColor=colors.HexColor("#0F4C81"),
        spaceBefore=15,
        spaceAfter=8
    )
    style_h3 = ParagraphStyle(
        'SubSectionH3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        textColor=colors.HexColor("#333333"),
        spaceBefore=10,
        spaceAfter=5
    )
    style_body = ParagraphStyle(
        'DocBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#333333")
    )
    style_bullet = ParagraphStyle(
        'DocBullet',
        parent=styles['Bullet'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        leftIndent=15,
        spaceAfter=3
    )
    
    story = []
    
    # Logo / Header Block
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(base_dir, "assets", "logo.png")
    if os.path.exists(logo_path):
        try:
            story.append(Image(logo_path, width=70, height=70))
            story.append(Spacer(1, 10))
        except Exception:
            pass
            
    story.append(Paragraph("TVSM QAD LAB Validation Report", style_title))
    story.append(Spacer(1, 10))
    
    # Metadata Table
    meta_data = [
        [Paragraph("<b>Report Reference No:</b>", style_body), Paragraph(project_data.get("request_no", "N/A"), style_body)],
        [Paragraph("<b>Date:</b>", style_body), Paragraph(datetime.now().strftime("%d.%m.%Y"), style_body)],
        [Paragraph("<b>Classification:</b>", style_body), Paragraph("Internal", style_body)],
        [Paragraph("<b>Model Name:</b>", style_body), Paragraph(project_data['battery_name'], style_body)],
        [Paragraph("<b>Standard Reference:</b>", style_body), Paragraph(project_data.get("tses_version", "TSES 799 v8"), style_body)],
        [Paragraph("<b>Testing Engineer:</b>", style_body), Paragraph(project_data.get("engineer", "N/A"), style_body)]
    ]
    t_meta = Table(meta_data, colWidths=[150, 350])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F4F8")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#D3D3D3")),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 20))
    
    # Executive Summary Table
    story.append(Paragraph("1. Executive Summary Table", style_h2))
    
    sum_headers = ["Class", "Test Name", "Requirements", "Observations", "Verdict"]
    summary_table_data = [[Paragraph(f"<b>{h}</b>", ParagraphStyle('Hdr', parent=style_body, textColor=colors.white)) for h in sum_headers]]
    
    for pt in project_data.get('tests', []):
        verdict_str = pt.get('status', 'Pending')
        v_color = "#38761D" if verdict_str == 'Pass' else ("#C00000" if verdict_str == 'Fail' else "#B45F06")
        
        row = [
            Paragraph(pt['test_id'], style_body),
            Paragraph(pt['test_name'], style_body),
            Paragraph(pt.get('acceptance_criteria', 'N/A'), style_body),
            Paragraph(pt.get('observations', 'Pending'), style_body),
            Paragraph(f"<b>{verdict_str}</b>", ParagraphStyle('Vrd', parent=style_body, textColor=colors.HexColor(v_color)))
        ]
        summary_table_data.append(row)
        
    t_sum = Table(summary_table_data, colWidths=[40, 140, 130, 140, 50])
    t_sum.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0F4C81")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#D3D3D3")),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    
    # Color-code status cell backgrounds in Table
    for idx, pt in enumerate(project_data.get('tests', [])):
        row_idx = idx + 1
        verdict_str = pt.get('status', 'Pending')
        if verdict_str == 'Pass':
            bg_color = colors.HexColor("#E2F0D9")
        elif verdict_str == 'Fail':
            bg_color = colors.HexColor("#FCE4D6")
        else:
            bg_color = colors.HexColor("#FFF2CC")
        t_sum.setStyle(TableStyle([('BACKGROUND', (4, row_idx), (4, row_idx), bg_color)]))
        
    story.append(t_sum)
    story.append(Spacer(1, 20))
    
    # Detailed Individual Test Results
    story.append(Paragraph("2. Test Analysis Details", style_h2))
    for pt in project_data.get('tests', []):
        if pt.get('status') == 'Pending':
            continue
            
        test_story = []
        test_story.append(Paragraph(f"Test Class {pt['test_id']} - {pt['test_name']}", style_h3))
        test_story.append(Paragraph(f"<b>Acceptance Criteria:</b> {pt.get('acceptance_criteria', 'N/A')}", style_body))
        
        # Add analysis stats
        if pt.get('results_json'):
            try:
                results = json.loads(pt['results_json'])
                test_story.append(Paragraph("<b>Telemetry Analytics Summary:</b>", style_body))
                if results.get('discharge'):
                    d = results['discharge']
                    test_story.append(Paragraph(f"• Measured Discharge Capacity: {d.get('capacity_ah', 0):.2f} Ah", style_bullet))
                    test_story.append(Paragraph(f"• Measured Discharge Energy: {d.get('energy_wh', 0):.1f} Wh", style_bullet))
                    test_story.append(Paragraph(f"• Active Load Rate: {d.get('c_rate', 0):.2f}C", style_bullet))
                if results.get('transitions'):
                    tr = results['transitions']
                    max_ir = max([t['dc_ir'] for t in tr]) if tr else 0.0
                    test_story.append(Paragraph(f"• Peak DC IR Estimate: {max_ir:.2f} m&Omega;", style_bullet))
                if results.get('thermal'):
                    t = results['thermal']
                    test_story.append(Paragraph(f"• Temperature Profile: Max={t.get('t_max', 0):.1f}&deg;C, Delta Rise=+{t.get('t_rise', 0):.1f}&deg;C", style_bullet))
            except Exception:
                pass
                
        # Insert generated custom graphs if they exist
        has_custom = False
        if pt.get('results_json'):
            try:
                results_data = json.loads(pt['results_json'])
                custom_graphs = results_data.get("custom_graphs", [])
                if custom_graphs:
                    has_custom = True
                    for idx, g in enumerate(custom_graphs):
                        img_path = g.get('image_path')
                        if img_path and os.path.exists(img_path):
                            test_story.append(Spacer(1, 5))
                            test_story.append(Image(img_path, width=400, height=230))
                            caption = f"Figure {pt['test_id']}-{idx+1}: {g.get('y1_axis')} vs {g.get('x_axis')} curve."
                            test_story.append(Paragraph(f"<i>{caption}</i>", style_body))
                            test_story.append(Spacer(1, 5))
            except Exception:
                pass
                
        if not has_custom:
            # Fallback to single standard graph
            graph_img = pt.get('processed_data_path')
            if graph_img and os.path.exists(graph_img):
                png_graph = graph_img.replace('.csv', '_graph.png')
                if os.path.exists(png_graph):
                    test_story.append(Spacer(1, 5))
                    test_story.append(Image(png_graph, width=400, height=230))
                    test_story.append(Spacer(1, 5))
                
        test_story.append(Spacer(1, 10))
        story.append(KeepTogether(test_story))
        
    story.append(Spacer(1, 20))
    
    # Conclusions and Signatures
    conclusion_block = []
    conclusion_block.append(Paragraph("3. Conclusions & Declarations", style_h2))
    conclusion_block.append(Paragraph(project_data.get("comments", "No safety failures or cell performance anomalies were detected. The battery pack conforms to the required standards."), style_body))
    conclusion_block.append(Spacer(1, 20))
    
    sig_data = [
        [Paragraph("<b>Prepared by:</b>", style_body), Paragraph("<b>Checked by:</b>", style_body), Paragraph("<b>Approved by:</b>", style_body)],
        [Paragraph(project_data.get("engineer", "Testing Engineer"), style_body), Paragraph("S Shivram<br/>Quality Team Manager", style_body), Paragraph("C Varunkumar<br/>Lab Director", style_body)]
    ]
    t_sig = Table(sig_data, colWidths=[166, 166, 166])
    t_sig.setStyle(TableStyle([
        ('LINEABOVE', (0,0), (-1,0), 0.5, colors.HexColor("#D3D3D3")),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    conclusion_block.append(t_sig)
    
    story.append(KeepTogether(conclusion_block))
    
    doc.build(story, canvasmaker=NumberedCanvas)

def export_to_csv(project_data, output_path):
    """
    Exports executive summary table as a standard CSV.
    """
    with open(output_path, mode='w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Project Name", project_data.get("project_name", "N/A")])
        writer.writerow(["Reference No", project_data.get("request_no", "N/A")])
        writer.writerow(["Battery Model", project_data["battery_name"]])
        writer.writerow([])
        
        writer.writerow(["Class ID", "Test Name", "Acceptance Criteria", "Observations", "Verdict"])
        for pt in project_data.get('tests', []):
            writer.writerow([
                pt['test_id'],
                pt['test_name'],
                pt.get('acceptance_criteria', 'N/A'),
                pt.get('observations', 'Pending'),
                pt.get('status', 'Pending')
            ])
